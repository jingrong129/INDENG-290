from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = PROJECT_ROOT / "Professor_Progress_Report.docx"


def add_metric_table(document: Document, title: str, df: pd.DataFrame) -> None:
    document.add_paragraph(title, style="Heading 2")
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if isinstance(value, float):
                cells[i].text = f"{value:.3f}"
            else:
                cells[i].text = str(value)
    document.add_paragraph("")


def main() -> None:
    benchmark = pd.read_csv(PROJECT_ROOT / "data/processed/benchmark_metrics.csv")
    bias_corrected = pd.read_csv(PROJECT_ROOT / "data/processed/bias_corrected_metrics.csv")
    cpi_brent_summary = pd.read_csv(PROJECT_ROOT / "data/processed/cpi_brent_summary.csv")
    rolling = pd.read_csv(PROJECT_ROOT / "data/processed/rolling_model_metrics.csv")
    panel = pd.read_csv(PROJECT_ROOT / "data/processed/eia_brent_forecasts.csv")
    evaluated = panel.loc[panel["horizon_years"].isin([3, 5]) & panel["actual_real_2025_usd_per_bbl"].notna()].copy()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Progress Report: Long-Term Oil Price Forecast Evaluation")
    run.bold = True
    run.font.size = Pt(14)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("UCB Energy Project").italic = True

    doc.add_paragraph(
        "This report summarizes the current status of our project on the reliability of long-term oil price forecasts for upstream investment planning. "
        "The objective is to evaluate whether widely used institutional forecasts are informative at planning horizons relevant for capital allocation, "
        "whether they exhibit systematic bias, and whether simple econometric models can improve upon them."
    )

    doc.add_paragraph("Research Question", style="Heading 1")
    doc.add_paragraph(
        "Our central question is whether long-horizon oil price forecasts provide decision-relevant information for upstream oil and gas firms. "
        "Following the project proposal, we focus on 3-year-ahead and 5-year-ahead forecasts of annual average Brent prices, because these horizons "
        "are closely aligned with project screening, development timing, and final investment decisions."
    )

    doc.add_paragraph("Data and Current Scope", style="Heading 1")
    doc.add_paragraph(
        "The current implementation uses EIA Annual Energy Outlook (AEO) Table 12 vintages as the core institutional forecast source. "
        "We extracted Brent Spot forecasts from publicly available AEO releases spanning 2013 to 2026, excluding 2024 because no corresponding AEO vintage exists. "
        "Realized Brent prices are taken from FRED series DCOILBRENTEU and aggregated to annual averages. "
        "To ensure comparability across vintages, all forecasts and realizations are converted into real 2025 U.S. dollars using CPIAUCSL. "
        "We also downloaded World Bank commodity forecast archives as supplementary material and incorporated Brent front-month futures (BZ=F) as a public market-information proxy for model extension."
    )
    doc.add_paragraph(
        f"In the current realized evaluation sample, we use {evaluated['vintage_year'].nunique()} EIA publication years and {len(evaluated.loc[evaluated['horizon_years'].isin([3,5])])} forecast observations across the 3-year and 5-year horizons."
    )

    doc.add_paragraph("Methodology", style="Heading 1")
    doc.add_paragraph(
        "We evaluate forecast performance using Mean Error (ME), Mean Percentage Error (MPE), Mean Absolute Error (MAE), Root Mean Squared Error (RMSE), and Mean Absolute Percentage Error (MAPE). "
        "ME is especially important because it captures the direction of bias in dollar terms, while MPE expresses that bias relative to the realized oil price level. "
        "As a benchmark, we use a martingale or no-change forecast equal to the most recently completed annual Brent price available at the time of release."
    )
    doc.add_paragraph(
        "To assess whether forecast accuracy can be improved, we also implemented two rolling out-of-sample models: "
        "(i) a rolling autoregressive specification based on lagged annual Brent prices, and "
        "(ii) a rolling AR model augmented with a market-information proxy derived from Brent front-month futures around the forecast release month."
    )

    doc.add_paragraph("Current Findings", style="Heading 1")
    doc.add_paragraph(
        "The main result to date is that EIA long-horizon Brent forecasts outperform the random-walk benchmark in RMSE at both the 3-year and 5-year horizons. "
        "This suggests that the EIA forecasts contain economically meaningful information and are not equivalent to a naive carry-forward rule."
    )
    doc.add_paragraph(
        "At the same time, EIA forecasts display a strong positive Mean Error in the realized sample. "
        "This indicates an optimistic bias: on average, projected Brent prices are above realized outcomes. "
        "From a capital budgeting perspective, this result is important because systematically optimistic oil price assumptions may translate into overstated project cash flows, inflated NPV and IRR calculations, and a tendency toward overinvestment."
    )

    add_metric_table(
        doc,
        "Benchmark Comparison",
        benchmark[["model", "horizon_years", "n_forecasts", "ME", "MPE", "MAE", "RMSE", "MAPE"]],
    )

    doc.add_paragraph(
        "To avoid using the EIA level forecast mechanically when historical optimism is present, we also estimated a bias-corrected EIA series. "
        "For each horizon, the correction uses only the mean error observed in earlier forecast vintages, so the exercise remains out of sample and does not use future information."
    )
    add_metric_table(
        doc,
        "Bias-Corrected EIA Comparison",
        bias_corrected[["model", "horizon_years", "n_forecasts", "ME", "MPE", "MAE", "RMSE", "MAPE"]],
    )

    doc.add_paragraph(
        "The rolling econometric models have not yet outperformed the EIA forecast. "
        "In the current specification, both the rolling AR model and the AR-plus-futures model produce larger RMSE values than the institutional forecast benchmark. "
        "This does not imply that model-based improvement is impossible; rather, it suggests that our current reduced-form implementation remains too simple relative to the information embedded in the EIA outlook."
    )

    add_metric_table(
        doc,
        "Rolling Model Comparison",
        rolling[["model", "horizon_years", "n_forecasts", "ME", "MAE", "RMSE", "MAPE"]],
    )

    doc.add_paragraph("Interpretation", style="Heading 1")
    doc.add_paragraph(
        "The evidence so far supports a nuanced conclusion. Institutional long-term forecasts appear useful in the sense that they outperform a naive benchmark, "
        "but they should not be taken at face value because their positive bias may distort investment decisions if used mechanically. "
        "A practical implication is that an upstream firm may want to treat institutional forecasts as informative baseline inputs while applying conservative adjustments or explicit bias corrections in project valuation."
    )
    doc.add_paragraph(
        "We also checked whether the CPI deflation step could be driving the result mechanically. "
        "The annual relationship between Brent and CPI is present, but the real-price construction remains transparent because the transformation is a direct scaling by observed CPI rather than a fitted model. "
        "The supplementary CPI-Brent summary table below is intended as a diagnostic on that point."
    )
    add_metric_table(doc, "CPI and Brent Diagnostic Summary", cpi_brent_summary)

    doc.add_paragraph("Limitations and Next Steps", style="Heading 1")
    doc.add_paragraph(
        "Three limitations remain. First, the main analysis currently centers on EIA because it provides the cleanest public vintage structure; "
        "IEA historical long-horizon Brent forecasts have not yet been integrated into a machine-readable panel. "
        "Second, the World Bank archive is supplementary because its oil series is defined as 'Crude oil, avg' rather than exact Brent. "
        "Third, the futures extension currently uses a front-month public proxy rather than a full long-dated Brent futures curve."
    )
    doc.add_paragraph(
        "The next stage of the project will focus on strengthening the institutional comparison, refining the market-information component, "
        "and developing a clearer bias-correction framework that can be translated into decision guidance for capital allocation."
    )

    fig1 = PROJECT_ROOT / "outputs/benchmark_metrics.png"
    fig2 = PROJECT_ROOT / "outputs/eia_vintages_vs_actual.png"
    if fig1.exists():
        doc.add_paragraph("Figure 1. Benchmark comparison", style="Heading 2")
        doc.add_picture(str(fig1), width=Inches(6.0))
    if fig2.exists():
        doc.add_paragraph("Figure 2. EIA forecast vintages versus realized Brent", style="Heading 2")
        doc.add_picture(str(fig2), width=Inches(6.0))
    fig3 = PROJECT_ROOT / "outputs/bias_correction_metrics.png"
    fig4 = PROJECT_ROOT / "outputs/cpi_brent_relationship.png"
    if fig3.exists():
        doc.add_paragraph("Figure 3. Bias-corrected EIA versus benchmark metrics", style="Heading 2")
        doc.add_picture(str(fig3), width=Inches(6.0))
    if fig4.exists():
        doc.add_paragraph("Figure 4. CPI and Brent diagnostic plots", style="Heading 2")
        doc.add_picture(str(fig4), width=Inches(6.0))

    doc.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
